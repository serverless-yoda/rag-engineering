# utils/generictext_utils.py

"""
Generic text extraction utilities for various file types.
This module provides functions to convert different input types
into clean, normalized text suitable for embedding and storage.
"""

import os
import json
import csv
import io
from pathlib import Path
from typing import Any, Dict, Optional

from html.parser import HTMLParser

# Optional deps listed in the environment toolset:
from pypdf import PdfReader
from docx import Document

# ---------- Helpers ----------

class _HTMLStripper(HTMLParser):
    """Simple HTML tag stripper using stdlib."""
    def __init__(self):
        super().__init__()
        self._chunks = []
    def handle_data(self, d):
        self._chunks.append(d)
    def get_text(self):
        return "".join(self._chunks)

def strip_html(html: str) -> str:
    s = _HTMLStripper()
    s.feed(html)
    return s.get_text()

TEXT_LIKE_MIMES = {
    "text/plain", "text/markdown", "text/csv",
    "application/json",
}

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME  = "application/pdf"

def _safe_read_utf8(path: Path) -> Optional[str]:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback: binary read then try decode with errors='replace'
        try:
            return path.read_bytes().decode("utf-8", errors="replace")
        except Exception:
            return None
    except Exception:
        return None

def _read_text_from_csv(path: Path) -> Optional[str]:
    try:
        with path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            rows = ["\t".join(row) for row in reader]
            return "\n".join(rows)
    except Exception:
        return None

def _read_text_from_json(path: Path) -> Optional[str]:
    try:
        obj = json.loads(path.read_text(encoding="utf-8"))
        # Pretty-print JSON deterministically
        return json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        try:
            # If not valid JSON, just return raw text
            return path.read_text(encoding="utf-8")
        except Exception:
            return None

def _read_text_from_pdf(path: Path) -> Optional[str]:
    try:
        with path.open("rb") as f:
            reader = PdfReader(f)
            parts = []
            for page in reader.pages:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    # Some pages may fail to extract text
                    continue
            text = "\n".join(p for p in parts if p)
            return text or None
    except Exception:
        return None

def _read_text_from_docx(path: Path) -> Optional[str]:
    try:
        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs]
        return "\n".join(paras).strip() or None
    except Exception:
        return None

def _read_text_from_fileobj(fileobj, name_hint: Optional[str] = None) -> Optional[str]:
    """
    Best-effort extraction from a file-like object. Reads into memory,
    so use carefully for large files.
    """
    try:
        # Try to peek content
        data = fileobj.read()
        if isinstance(data, bytes):
            # Guess from name hint
            if name_hint and name_hint.lower().endswith(".pdf"):
                try:
                    reader = PdfReader(io.BytesIO(data))
                    parts = []
                    for page in reader.pages:
                        try:
                            parts.append(page.extract_text() or "")
                        except Exception:
                            continue
                    text = "\n".join(p for p in parts if p)
                    return text or None
                except Exception:
                    pass
            if name_hint and name_hint.lower().endswith(".docx"):
                try:
                    bio = io.BytesIO(data)
                    doc = Document(bio)  # python-docx can read file-like
                    paras = [p.text for p in doc.paragraphs]
                    return "\n".join(paras).strip() or None
                except Exception:
                    pass
            # Try UTF-8 decode as generic text
            try:
                return data.decode("utf-8", errors="replace")
            except Exception:
                return None
        else:
            # If .read() returned str
            if isinstance(data, str):
                return data
            return None
    except Exception:
        return None

def _ensure_normalized(item: Any) -> Dict[str, Any]:
    """
    Accepts either a normalized dict (from normalize_items) or raw input,
    and returns a normalized dict with keys: name, mime_type, source, content?
    """
    # If it already looks normalized (has mime_type and source)
    if isinstance(item, dict) and "mime_type" in item and "source" in item:
        return item

    # Otherwise, reuse the normalize_single_item logic inline (simplified),
    # or better: call your normalize_items([item])[0]. Here we do a minimal version.
    from pathlib import Path as _Path
    import mimetypes as _mimetypes

    def _guess_mime_from_name(name: Optional[str]) -> str:
        if not name:
            return "application/octet-stream"
        mime, _ = _mimetypes.guess_type(name)
        return mime or "application/octet-stream"

    # dict raw (not normalized)
    if isinstance(item, dict):
        name = item.get("name") or item.get("filename") or "unnamed"
        mime = item.get("mime_type") or item.get("mime") or _guess_mime_from_name(name)
        norm = {"name": name, "mime_type": mime, "source": {"type": "dict", "value": item}}
        if "content" in item:
            norm["content"] = item["content"]
        return norm

    # bytes-like
    if isinstance(item, (bytes, bytearray, memoryview)):
        return {"name": "bytes", "mime_type": "application/octet-stream", "source": {"type": "bytes", "value": bytes(item)}}

    # str: could be a path or raw text
    if isinstance(item, str):
        p = _Path(item)
        if p.exists() and p.is_file():
            mime = _guess_mime_from_name(p.name)
            return {"name": p.name, "mime_type": mime, "source": {"type": "path", "value": str(p)}}
        else:
            return {"name": "text", "mime_type": "text/plain", "source": {"type": "bytes", "value": item.encode("utf-8")}, "content": item}

    # Path
    if isinstance(item, _Path):
        mime = _guess_mime_from_name(item.name)
        return {"name": item.name, "mime_type": mime, "source": {"type": "path", "value": str(item)}}

    # file-like
    if hasattr(item, "read") and callable(getattr(item, "read")):
        name = getattr(item, "name", None) or "fileobj"
        mime = _guess_mime_from_name(name if isinstance(name, str) else None)
        return {"name": os.path.basename(name) if isinstance(name, str) else "fileobj", "mime_type": mime, "source": {"type": "fileobj", "value": item}}

    # Fallback
    return {"name": "unknown", "mime_type": "application/octet-stream", "source": {"type": "unknown", "value": item}}

# ---------- Main: to_text_content ----------

def file_to_text_content(item: Any) -> Optional[str]:
    """
    Convert a normalized item (or raw input) to clean text.
    Handles:
      - text/plain, text/markdown, text/csv, application/json
      - application/pdf (PyPDF2 extraction)
      - application/vnd.openxmlformats-officedocument.wordprocessingml.document (python-docx)
      - HTML stripping if mime suggests text/html or content looks like HTML
      - bytes: UTF-8 decode (replace errors)
      - paths: reads and extracts per mime
    Returns None if no text can be reasonably produced.
    """
    norm = _ensure_normalized(item)

    # If content already present and is text-like, use it
    content = norm.get("content")
    if isinstance(content, bytes):
        try:
            content = content.decode("utf-8", errors="replace")
        except Exception:
            content = None
    if isinstance(content, str) and content:
        # If it smells like HTML, strip tags
        mt = norm.get("mime_type", "")
        if mt == "text/html" or ("<html" in content.lower() or "<body" in content.lower()):
            return strip_html(content)
        return content

    # Dispatch by source type
    source = norm.get("source", {})
    src_type = source.get("type")
    mime = norm.get("mime_type", "application/octet-stream")

    # ---- bytes / fileobj ----
    if src_type == "bytes":
        data = source.get("value", b"")
        if isinstance(data, bytes):
            # Try PDF/DOCX decode when name gives a hint
            name_hint = norm.get("name")
            text = None
            if name_hint and name_hint.lower().endswith(".pdf"):
                try:
                    reader = PdfReader(io.BytesIO(data))
                    parts = []
                    for page in reader.pages:
                        try:
                            parts.append(page.extract_text() or "")
                        except Exception:
                            continue
                    text = "\n".join(p for p in parts if p) or None
                except Exception:
                    text = None
            elif name_hint and name_hint.lower().endswith(".docx"):
                try:
                    doc = Document(io.BytesIO(data))
                    paras = [p.text for p in doc.paragraphs]
                    text = "\n".join(paras).strip() or None
                except Exception:
                    text = None
            # Fallback UTF-8 decode
            if not text:
                try:
                    text = data.decode("utf-8", errors="replace")
                except Exception:
                    text = None
            # Strip HTML if needed
            if isinstance(text, str) and (mime == "text/html" or ("<html" in text.lower())):
                text = strip_html(text)
            return text

    if src_type == "fileobj":
        fileobj = source.get("value")
        name_hint = norm.get("name")
        text = _read_text_from_fileobj(fileobj, name_hint=name_hint)
        # Strip HTML if needed
        if isinstance(text, str) and (mime == "text/html" or ("<html" in text.lower())):
            text = strip_html(text)
        return text

    # ---- path ----
    if src_type == "path":
        p = Path(source.get("value"))
        if not p.exists() or not p.is_file():
            return None

        # Handle common text-like mimes
        if mime in TEXT_LIKE_MIMES or p.suffix.lower() in {".txt", ".md"}:
            text = _safe_read_utf8(p)
            if not text:
                return None
            # HTML strip if necessary
            if mime == "text/html" or ("<html" in text.lower()):
                return strip_html(text)
            if mime == "application/json" or p.suffix.lower() == ".json":
                # Normalize/pretty print
                return _read_text_from_json(p) or text
            if mime == "text/csv" or p.suffix.lower() == ".csv":
                return _read_text_from_csv(p) or text
            return text

        # PDF
        if mime == PDF_MIME or p.suffix.lower() == ".pdf":
            return _read_text_from_pdf(p)

        # DOCX
        if mime == DOCX_MIME or p.suffix.lower() == ".docx":
            return _read_text_from_docx(p)

        # Unknown binary -> no text
        return None

    # ---- dict / unknown ----
    if src_type == "dict":
        # Try common fields
        v = source.get("value", {})
        # If it contains 'text' or similar
        for key in ("text", "message", "body", "content"):
            if isinstance(v.get(key), str) and v.get(key).strip():
                t = v.get(key)
                if mime == "text/html" or ("<html" in t.lower()):
                    return strip_html(t)
                return t
        # If dict itself is JSON-like, pretty print
        try:
            return json.dumps(v, ensure_ascii=False, indent=2)
        except Exception:
            return None

    # Fallback
    return None
